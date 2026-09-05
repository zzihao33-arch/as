from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("deliverables/TYG-面单推送接口技术评审稿-v0.9.docx")

# compact_reference_guide token map, with the named customer-pack header override.
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
AMBER = "FFF4D6"
RED_FILL = "FDEBEC"
RED = "9B1C1C"
GREEN = "1F3A5F"


def set_run_font(run, *, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in kwargs.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths, *, header=True, fill=LIGHT_BLUE, indent=TABLE_INDENT_DXA):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    # python-docx creates equal grid columns by default. Replace them so the
    # OOXML table grid matches the explicit cell widths in every row.
    tbl_grid = table._tbl.tblGrid
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = borders.find(qn(f"w:{edge}"))
        if item is None:
            item = OxmlElement(f"w:{edge}")
            borders.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "4")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), "C9D2DD")
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_index])
            set_cell_margins(cell, **CELL_MARGINS)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.1
            if header and row_index == 0:
                set_cell_shading(cell, fill)
                tr_pr = row._tr.get_or_add_trPr()
                if tr_pr.find(qn("w:tblHeader")) is None:
                    header_el = OxmlElement("w:tblHeader")
                    header_el.set(qn("w:val"), "true")
                    tr_pr.append(header_el)


def set_table_text(cell, text, *, bold=False, color="000000", size=9.4, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.clear()
    r = p.add_run(str(text))
    set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_table(doc, headers, rows, widths, *, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, value in enumerate(headers):
        set_table_text(table.rows[0].cells[idx], value, bold=True, color=NAVY, size=9.2)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_table_text(cells[idx], value, size=font_size)
    configure_table(table, widths)
    return table


def paragraph_border_bottom(paragraph, color="D7DBE2", size="6", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_repeatable_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("CM-HUB × TYG  |  面单推送接口技术评审稿")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    paragraph_border_bottom(p, color="D7DBE2", size="4", space="4")
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(3)
    r = fp.add_run("内部技术评审稿  |  第 ")
    set_run_font(r, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    fp._p.append(field)
    r = fp.add_run(" 页")
    set_run_font(r, size=8.5, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    # Remove the empty run that some templates place in the heading paragraph.
    for run in p.runs[:-1]:
        run.text = ""
    return p


def add_body(doc, text, *, bold_prefix=None, color="000000", after=6):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, color=color, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=11, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.clear()
    r = p.add_run(text)
    set_run_font(r, name="Consolas", size=8.2, color="1F2937")
    configure_table(table, [9360], header=False)
    return table


def add_callout(doc, title, body, *, fill=CALLOUT, title_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=150, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=title_color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r = p2.add_run(body)
    set_run_font(r, size=10, color="2D3748")
    configure_table(table, [9360], header=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def set_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("客户系统对接 | 技术评审材料")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("CM-HUB × TYG")
    set_run_font(r, size=27, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("面单推送接口技术评审稿")
    set_run_font(r, size=22, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("三字段 JSON / PDF Base64 一体推送方案")
    set_run_font(r, size=13, color=MUTED)

    metadata = doc.add_table(rows=4, cols=2)
    meta_rows = [
        ("文档版本", "v0.9（内部技术评审）"),
        ("更新日期", "2026-08-31"),
        ("适用对象", "CM-HUB 产品、后端、运维与安全评审人员"),
        ("评审目的", "确认 TYG 所需三字段 Base64 推送模型及上线条件"),
    ]
    for row, values in zip(metadata.rows, meta_rows):
        set_table_text(row.cells[0], values[0], bold=True, color=NAVY, size=9.5)
        set_table_text(row.cells[1], values[1], size=9.5)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
    configure_table(metadata, [2160, 7200], header=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_callout(
        doc,
        "评审状态：尚未对外发布",
        "本文定义的是拟新增的“单票三字段 Base64”接口合同。当前生产系统已实现的流程为“批次 JSON + 原始 PDF 上传”，与本稿不同；未完成接口开发、压测、鉴权配置和联调验收前，不得将本文作为生产接口承诺发送给客户。",
        fill=RED_FILL,
        title_color=RED,
    )

    add_heading(doc, "一、评审结论（供决策）", 1)
    add_body(doc, "建议批准新增一个独立的单票面单推送接口，以匹配 TYG “原单号、转单号、PDF Base64” 三字段一体推送的诉求。接口应采用异步受理、持久化幂等和受控并发，避免数万单日量在高峰期造成同步处理拥塞。")
    add_table(
        doc,
        ["评审项", "建议结论", "上线前状态"],
        [
            ("请求模型", "POST JSON；正文仅包含 3 个业务字段", "待后端确认"),
            ("受理方式", "完成校验与持久化后返回 202 Accepted；后续异步处理", "待实现"),
            ("身份认证", "沿用客户专属 X-API-Key；凭据仅通过安全渠道交付", "现有能力可复用"),
            ("容量基线", "日均数万票；默认 600 请求/分钟/Key，起始并发 5", "待压测验证"),
            ("对外文档", "内部评审通过后固化为客户版 v1.0", "待发布"),
        ],
        [1850, 5000, 2510],
        font_size=8.9,
    )


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_styles(doc)
    set_repeatable_header(section)
    add_cover(doc)

    doc.add_page_break()
    add_heading(doc, "二、背景、边界与现状差异", 1)
    add_body(doc, "TYG 提出的对接方式是：每次推送一票业务数据和对应快递面单，接口正文只接收原单号、转单号和 PDF 格式面单的 Base64 字符串。该方式可减少客户侧的调用步骤，但会使每个 JSON 请求体携带二进制文件编码后的内容。")
    add_table(
        doc,
        ["项目", "当前已实现能力", "本评审稿拟新增能力"],
        [
            ("数据写入", "按批次写入空提预报与物流单据", "按单票写入原单号与转单号映射"),
            ("面单上传", "单独上传原始 application/pdf 文件", "在 JSON 正文携带 PDF Base64"),
            ("成功语义", "同步写入后返回成功", "持久化受理后返回 202，后续异步处理"),
            ("字段模型", "批次、提单、物流单据等完整上下文", "客户要求的 3 个业务字段；不要求客户补充批次信息"),
        ],
        [1950, 3700, 3710],
        font_size=9.0,
    )
    add_callout(doc, "范围控制", "本次仅讨论 TYG 向 CM-HUB 单向推送单票映射与 PDF 面单。不包含仓库状态回写、打印事件回传、面单修改/覆盖流程或对账报表接口。", fill=AMBER, title_color="7A5A00")

    add_heading(doc, "三、拟定接口合同", 1)
    add_heading(doc, "3.1 资源与调用方式", 2)
    add_code_block(doc, "POST https://api.cmhubtool.com/api/v1/label-pushes\nContent-Type: application/json\nX-API-Key: cmh_live_<由 CM-HUB 通过安全渠道提供>\nIdempotency-Key: <稳定且唯一的请求键>\nX-Request-ID: <可选，建议每次请求唯一>")
    add_body(doc, "说明：路径 /api/v1/label-pushes 为本评审稿建议路径，需由后端实现后才可正式启用。客户正文不需要客户 ID；系统根据 X-API-Key 识别推送方并绑定数据归属。")

    add_heading(doc, "3.2 请求头", 2)
    add_table(
        doc,
        ["Header", "必填", "规则与用途"],
        [
            ("X-API-Key", "是", "TYG 专属服务端 API Key；不得置于 URL、浏览器代码或普通日志。"),
            ("Content-Type", "是", "固定为 application/json。"),
            ("Idempotency-Key", "是", "8–128 位字母、数字、下划线或连字符；超时重试必须复用。"),
            ("X-Request-ID", "否", "8–64 位字母、数字、下划线或连字符；用于双方排障。"),
        ],
        [2050, 850, 6460],
        font_size=9.0,
    )

    add_heading(doc, "3.3 请求正文（严格限定为三个业务字段）", 2)
    add_table(
        doc,
        ["字段", "类型", "必填", "规则"],
        [
            ("originalTrackingNo", "string", "是", "原单号。去首尾空格后 1–128 字符；作为业务唯一键的一部分。"),
            ("transferTrackingNo", "string", "是", "转单号/快递单号。去首尾空格后 1–128 字符；作为业务唯一键的一部分。"),
            ("labelPdfBase64", "string", "是", "PDF 原始字节的标准 Base64 编码；不得包含 data:application/pdf;base64, 前缀、换行或 URL-safe Base64 变体。"),
        ],
        [2200, 900, 700, 5560],
        font_size=8.85,
    )
    add_code_block(doc, "{\n  \"originalTrackingNo\": \"HHWV06218005702YQ\",\n  \"transferTrackingNo\": \"9400111899560000000000\",\n  \"labelPdfBase64\": \"JVBERi0xLjcKJc...（完整 PDF Base64 内容）\"\n}")
    add_callout(doc, "PDF 校验", "服务端必须先完成 Base64 解码，再校验 PDF 文件头、文件尾和原始字节大小；不能仅依据字段名或文件扩展名判断。服务端自动计算 SHA-256 用于审计、去重与存储，不要求 TYG 在正文额外传递哈希字段。")

    add_heading(doc, "3.4 容量与报文限制（建议默认值）", 2)
    add_table(
        doc,
        ["项目", "建议默认值", "说明"],
        [
            ("原始 PDF 最大大小", "5 MiB", "Base64 编码后的请求体约增加 33%；超过限制直接返回 413。"),
            ("请求体最大大小", "7 MiB", "包含 JSON 字段及 Base64 内容；网关、应用与反向代理需统一配置。"),
            ("单票处理模型", "异步受理", "完成验证与可靠持久化后返回 202，后续仓储/打印链路异步处理。"),
            ("限流", "600 请求/分钟/Key", "高峰期需按实际压测结果调整；返回 429 时客户应退避重试。"),
            ("客户端起始并发", "5", "可依据 429、超时率和 p95 响应时间逐步调优，禁止瞬间建立数千连接。"),
        ],
        [2150, 1500, 5710],
        font_size=8.9,
    )
    add_body(doc, "容量说明：以 5 万票/日、每份原始 PDF 平均 100 KiB 估算，原始文件约 4.8 GiB/日，Base64 网络传输约 6.4 GiB/日。容量设计必须按高峰小时而非日均值预留，并监控请求大小、队列积压、对象存储写入与失败重试。")

    doc.add_page_break()
    add_heading(doc, "四、受理、幂等与重复推送规则", 1)
    add_heading(doc, "4.1 成功受理", 2)
    add_body(doc, "接口仅在请求通过基础格式校验、PDF 校验且请求数据已可靠持久化后返回成功。202 表示 CM-HUB 已接受该票请求，不表示仓库已打印或已完成后续业务操作。")
    add_code_block(doc, "HTTP/1.1 202 Accepted\n{\n  \"data\": {\n    \"originalTrackingNo\": \"HHWV06218005702YQ\",\n    \"transferTrackingNo\": \"9400111899560000000000\",\n    \"receiptNo\": \"lbl_01J7TYG000001\",\n    \"status\": \"ACCEPTED\",\n    \"duplicate\": false\n  },\n  \"requestId\": \"TYG-20260831-000001\"\n}")
    add_heading(doc, "4.2 幂等规则", 2)
    add_body(doc, "建议业务唯一键为 originalTrackingNo + transferTrackingNo。Idempotency-Key 用于识别同一业务操作的网络重试，不能用同一个键发送不同的请求内容。建议键格式：TYG-<原单号>-<转单号>-v1；如长度可能超过 128 位，使用双方约定的稳定哈希缩写。")
    add_table(
        doc,
        ["情形", "系统行为", "客户动作"],
        [
            ("同一幂等键 + 完全相同请求", "返回原受理结果，并标记 duplicate: true 或 idempotentReplay: true。", "视为成功；停止重试。"),
            ("同一幂等键 + 不同请求", "返回 409 IDEMPOTENCY_CONFLICT。", "不可盲目重试；使用新的业务版本键并联系 CM-HUB。"),
            ("同一单号组合 + 相同 PDF", "复用已存储面单，不生成重复资产。", "视为成功。"),
            ("同一单号组合 + 不同 PDF", "默认返回 409 LABEL_ALREADY_EXISTS，禁止静默覆盖。", "走双方确认的更换面单流程；本期不在客户接口范围内。"),
        ],
        [2550, 4300, 2510],
        font_size=8.85,
    )

    add_heading(doc, "4.3 服务端处理流程", 2)
    add_code_block(doc, "TYG 服务端\n  → 身份认证、限流、幂等检查\n  → Base64 解码与 PDF 结构校验\n  → 业务映射与面单资产可靠持久化\n  → 返回 202 Accepted + requestId / receiptNo\n  → 异步队列驱动后续仓储、打印可用性与审计事件")

    add_heading(doc, "五、失败响应与重试策略", 1)
    add_table(
        doc,
        ["HTTP", "错误码示例", "客户处理方式"],
        [
            ("400", "VALIDATION_ERROR / INVALID_BASE64", "修正字段或编码；不要原样自动重试。"),
            ("401", "INVALID_API_KEY", "检查凭据；疑似泄露时联系 CM-HUB 撤销并换发。"),
            ("403", "INSUFFICIENT_SCOPE", "联系 CM-HUB 配置权限。"),
            ("409", "IDEMPOTENCY_CONFLICT / LABEL_ALREADY_EXISTS", "停止自动重试，按错误说明人工核查。"),
            ("413", "PAYLOAD_TOO_LARGE", "压缩或拆分原始 PDF；不得切分同一份 PDF 的 Base64 字符串。"),
            ("415", "UNSUPPORTED_MEDIA_TYPE", "改为 Content-Type: application/json。"),
            ("422", "INVALID_LABEL_PDF", "确认解码后文件是完整 PDF，且未包含图片或 HTML 错误页。"),
            ("429", "RATE_LIMITED", "降低并发，读取 Retry-After；带随机抖动指数退避。"),
            ("500 / 503", "INTERNAL_ERROR / TEMPORARILY_UNAVAILABLE", "使用原 Idempotency-Key 重试；最大 5 次后告警。"),
        ],
        [650, 3380, 5330],
        font_size=8.6,
    )
    add_body(doc, "推荐退避间隔为 1 秒、2 秒、4 秒、8 秒、16 秒，并在每次间隔上加入随机抖动。对于 HTTP 400、401、403、413、415、422 和大多数 409，不应进行原样自动重试。若客户端因网络超时未获得响应，也必须使用相同的 Idempotency-Key 再次调用。")

    doc.add_page_break()
    add_heading(doc, "六、安全、数据保护与运维要求", 1)
    add_bullet(doc, "仅允许 TYG 服务端通过 HTTPS 调用；禁止浏览器、移动端或公开脚本直接调用。")
    add_bullet(doc, "X-API-Key 必须存放在密钥管理系统或服务端环境变量中；不得进入 URL、截图、前端包或普通应用日志。")
    add_bullet(doc, "访问日志不得写入完整 labelPdfBase64、完整 API Key、收件人地址或其他个人信息；排障仅使用 requestId、receiptNo、单号脱敏值和哈希。")
    add_bullet(doc, "面单文件应采用私有对象存储、最小权限访问和可追溯的下载审计；保留期限由双方业务与合规要求确认。")
    add_bullet(doc, "网关、反向代理和应用层必须使用一致的请求体上限与超时策略，避免边缘层允许而应用层拒绝的配置偏差。")
    add_bullet(doc, "API Key 支持轮换、停用和最小权限；如双方有固定出口 IP，可在后续启用 IP 白名单作为附加防护。")

    add_heading(doc, "七、压测与联调验收", 1)
    add_body(doc, "在对外发布客户版 v1.0 前，至少完成以下验收。所有测试数据必须使用脱敏单号和非生产面单。")
    acceptance = [
        "完成一票正常请求，确认 202、receiptNo、存储资产和审计事件均可追溯。",
        "用相同 Idempotency-Key 重放相同请求，确认不产生重复资产。",
        "用相同 Idempotency-Key 提交不同内容，确认返回 409 IDEMPOTENCY_CONFLICT。",
        "提交非法 Base64、非 PDF 内容、超尺寸 PDF 和缺失字段，确认得到明确错误码且不落脏数据。",
        "在受控并发下完成不少于 1,000 票压测；记录 p50/p95 响应、错误率、队列积压、存储写入和 CPU/内存。",
        "按预计峰值的 1.5 倍进行容量评估，验证 429、退避与恢复行为。",
        "完成 API Key 轮换、权限不足、日志脱敏和异常告警演练。",
        "提供测试环境、技术联系人、故障反馈路径和切换窗口后，再与 TYG 进行正式联调。",
    ]
    for item in acceptance:
        add_number(doc, item)

    add_heading(doc, "八、上线前待决事项", 1)
    add_table(
        doc,
        ["待决事项", "建议", "责任方", "状态"],
        [
            ("接口是否新增", "批准 /api/v1/label-pushes，并保留现有批次+原始 PDF 接口兼容。", "产品 / 后端", "待决"),
            ("报文上限", "默认原始 PDF 5 MiB、请求体 7 MiB；根据真实样本校准。", "后端 / 运维", "待压测"),
            ("异步落地语义", "仅在业务映射和面单资产可靠持久化后返回 202。", "后端", "待设计"),
            ("重复面单策略", "相同文件幂等复用；不同文件禁止静默覆盖。", "产品 / 业务", "待确认"),
            ("保留与访问", "明确面单保留期限、私有存储、审计和删除策略。", "业务 / 安全", "待确认"),
            ("客户版发布", "所有上述事项关闭后，将本文去除内部说明并发布客户版 v1.0。", "项目负责人", "待发布"),
        ],
        [1900, 4050, 1800, 1610],
        font_size=8.55,
    )

    add_heading(doc, "附录：建议给 TYG 的对接说明（评审通过后使用）", 1)
    add_callout(
        doc,
        "对客户说明稿",
        "我们将提供单票面单推送接口。每次请求仅需传原单号、转单号及 PDF 面单的 Base64 内容；接口采用异步受理，完成基础校验和可靠保存后返回受理结果。请使用服务端 API Key 调用，并按接口文档提供的幂等键、文件大小和重试规则执行。正式环境地址及凭据将通过安全渠道提供。",
        fill=LIGHT_BLUE,
        title_color=NAVY,
    )

    props = doc.core_properties
    props.title = "CM-HUB × TYG 面单推送接口技术评审稿 v0.9"
    props.subject = "PDF Base64 三字段一体推送接口评审"
    props.author = "CM-HUB"
    props.keywords = "TYG, API, PDF, Base64, 面单, 技术评审"
    doc.save(OUT)


if __name__ == "__main__":
    build()
