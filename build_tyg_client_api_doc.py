from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_tyg_api_review_doc import (
    AMBER,
    BLUE,
    CALLOUT,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MUTED,
    NAVY,
    add_body,
    add_bullet,
    add_callout,
    add_code_block,
    add_heading,
    add_number,
    add_table,
    configure_table,
    paragraph_border_bottom,
    set_cell_shading,
    set_repeatable_header,
    set_run_font,
    set_styles,
    set_table_text,
)


OUT = Path("deliverables/TYG-系统对接API接口文档-v1.1-客户评审稿.docx")


def set_customer_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("CM-HUB × TYG  |  系统对接 API 接口文档")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    paragraph_border_bottom(p, color="D7DBE2", size="4", space="4")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(3)
    r = fp.add_run("客户评审稿 v1.1  |  第 ")
    set_run_font(r, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    fp._p.append(field)
    r = fp.add_run(" 页")
    set_run_font(r, size=8.5, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("客户系统对接 | API CONTRACT")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("CM-HUB × TYG")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run("系统对接 API 接口文档")
    set_run_font(r, size=23, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("空运提单预报 · 统一面单推送 · 版本化覆盖")
    set_run_font(r, size=13, color=MUTED)

    metadata = doc.add_table(rows=5, cols=2)
    rows = [
        ("文档版本", "v1.1（客户评审稿）"),
        ("更新日期", "2026-08-31"),
        ("接口状态", "待开发部署 / 待联调"),
        ("适用系统", "TYG 服务端 → CM-HUB"),
        ("生产基址", "https://api.cmhubtool.com/api/v1（部署后启用）"),
    ]
    for row, values in zip(metadata.rows, rows):
        set_table_text(row.cells[0], values[0], bold=True, color=NAVY, size=9.5)
        set_table_text(row.cells[1], values[1], size=9.5)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
    configure_table(metadata, [2160, 7200], header=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_callout(
        doc,
        "文档说明",
        "本文用于双方确认接口合同。当前仅为客户评审稿，接口尚未开发部署；测试环境、生产环境和 API Key 将在部署后通过安全渠道提供，未取得联调凭据前不可直接调用。",
        fill=LIGHT_BLUE,
        title_color=NAVY,
    )


def add_common_response(doc):
    add_heading(doc, "3.3 通用成功响应", 2)
    add_code_block(doc, "{\n  \"code\": \"SUCCESS\",\n  \"message\": \"接收成功\",\n  \"data\": {},\n  \"requestId\": \"TYG-20260831-000001\"\n}")
    add_body(doc, "requestId 是双方排查问题的唯一请求编号。TYG 应保存该值；排障时无需发送完整 PDF Base64。")

    add_heading(doc, "3.4 通用失败响应", 2)
    add_code_block(doc, "{\n  \"code\": \"INVALID_LABEL_PDF\",\n  \"message\": \"labelBase64 解码后不是完整的 PDF 文件\",\n  \"requestId\": \"TYG-20260831-000002\"\n}")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_styles(doc)
    set_customer_header(section)
    add_cover(doc)

    add_heading(doc, "一、对接目标", 1)
    add_body(doc, "TYG 先向 CM-HUB 推送空运提单预报数据，再通过同一个接口逐票推送该提单下的原单号、转单号和 PDF 面单。首次上传、迟到补传、同单号换 PDF 以及转单号变更均使用统一面单接口；CM-HUB 按空运提单管理预报与实收差异、面单版本和仓库打印。")
    add_table(
        doc,
        ["顺序", "接口", "用途"],
        [
            ("1", "POST /api/v1/air-shipments", "创建或更新空运提单预报。"),
            ("2", "POST /api/v1/label-pushes", "统一处理首次上传、补传、换 PDF 和转单号更新。"),
        ],
        [780, 3500, 5080],
        font_size=9.1,
    )
    add_code_block(doc, "空运提单预报\n  → 统一接口逐票并发推送面单\n  → CM-HUB 可靠保存后逐票返回成功\n  → 迟到包裹继续补传；新版 PDF 或新转单号仍调用同一接口\n  → 仓库始终查看和打印最新有效版本")

    add_heading(doc, "二、业务规则摘要", 1)
    rules = [
        "必须先成功创建空运提单预报，再推送该提单下的面单；找不到提单时不暂存孤立面单。",
        "TYG 提供的是预报箱数、预报包裹数和预报重量；仓库实收数据由 CM-HUB 内部记录。",
        "仓库开始收货前允许 TYG 更新预报；开始收货后锁定预报字段，但仍允许新增、补传和更新面单。",
        "空运提单到仓或业务关闭后，统一面单接口仍持续接收首次面单、迟到补传和新版面单；系统记录关闭后补推标识。",
        "面单数量超过预报时仍接收有效数据，并在 CM-HUB 中显示差异和异常提醒。",
        "有效包裹数按唯一原单号统计；重复请求、换 PDF 和转单号更新都不增加包裹数。",
        "同一原单号可在原空运提单内版本化更新转单号，但不能自动改到另一张空运提单；新转单号不得已绑定其他原单号。",
    ]
    for rule in rules:
        add_bullet(doc, rule)

    doc.add_page_break()
    add_heading(doc, "三、通用协议", 1)
    add_heading(doc, "3.1 环境与传输", 2)
    add_table(
        doc,
        ["项目", "约定"],
        [
            ("协议", "仅支持 HTTPS。"),
            ("数据格式", "JSON，UTF-8 编码。"),
            ("测试环境", "接口部署后另行提供。"),
            ("生产基址", "https://api.cmhubtool.com/api/v1"),
            ("调用方", "仅允许 TYG 服务端调用，不允许浏览器或移动端直接调用。"),
        ],
        [1900, 7460],
        font_size=9.2,
    )

    add_heading(doc, "3.2 认证与请求头", 2)
    add_code_block(doc, "Content-Type: application/json\nX-API-Key: <TYG 专属 API Key>\nIdempotency-Key: <本次业务提交的稳定唯一键>\nX-Request-ID: <可选，建议每次请求唯一>")
    add_table(
        doc,
        ["Header", "必填", "说明"],
        [
            ("Content-Type", "是", "固定为 application/json。"),
            ("X-API-Key", "是", "TYG 专属服务器端密钥；测试与生产使用不同密钥。"),
            ("Idempotency-Key", "是", "8–128 位；真实业务变更使用新键，网络重试必须复用原键。"),
            ("X-Request-ID", "否", "8–64 位字母、数字、下划线或连字符；用于链路追踪。"),
        ],
        [2100, 850, 6410],
        font_size=9.0,
    )
    add_body(doc, "API Key 不得放入 URL、网页代码、移动端、截图或普通日志。若 TYG 具备固定出口公网 IP，双方可在联调后增加 IP 白名单。")
    add_body(doc, "同一 Idempotency-Key 和完全相同请求体重复提交时返回原结果；同一键对应不同请求体时返回 IDEMPOTENCY_CONFLICT。首次上传、换 PDF 或更换转单号属于不同业务提交，必须分别使用新键；旧请求重试必须继续使用旧键，不得换新键重发。")
    add_common_response(doc)

    doc.add_page_break()
    add_heading(doc, "四、接口一：空运提单预报", 1)
    add_code_block(doc, "POST /api/v1/air-shipments")
    add_body(doc, "用于创建或更新一张空运提单的预报数据。TYG 必须取得本接口成功响应后，才能推送该提单下的逐票面单。")

    add_heading(doc, "4.1 请求字段", 2)
    add_table(
        doc,
        ["字段", "类型", "必填", "规则"],
        [
            ("airWaybillNo", "string", "是", "空运提单号；去除首尾空格后最长 32 字符。"),
            ("forecastCartons", "integer", "是", "预报总箱数，正整数。"),
            ("forecastPackages", "integer", "是", "预报总包裹数，正整数。"),
            ("forecastWeight", "number", "是", "预报总重量，大于 0，最多 3 位小数。"),
            ("weightUnit", "string", "是", "重量单位，仅支持 KG 或 LB。"),
        ],
        [2350, 1000, 700, 5310],
        font_size=8.9,
    )

    add_heading(doc, "4.2 请求示例", 2)
    add_code_block(doc, "{\n  \"airWaybillNo\": \"180-98109734\",\n  \"forecastCartons\": 120,\n  \"forecastPackages\": 50000,\n  \"forecastWeight\": 1850.5,\n  \"weightUnit\": \"KG\"\n}")

    add_heading(doc, "4.3 成功响应", 2)
    add_code_block(doc, "{\n  \"code\": \"SUCCESS\",\n  \"message\": \"空运提单预报保存成功\",\n  \"data\": {\n    \"airWaybillNo\": \"180-98109734\",\n    \"forecastCartons\": 120,\n    \"forecastPackages\": 50000,\n    \"forecastWeight\": 1850.5,\n    \"weightUnit\": \"KG\",\n    \"duplicate\": false,\n    \"updated\": false\n  },\n  \"requestId\": \"TYG-20260831-000101\"\n}")

    add_heading(doc, "4.4 创建、更新与锁定", 2)
    add_table(
        doc,
        ["情形", "系统处理"],
        [
            ("首次推送提单号", "创建预报并返回成功。"),
            ("完全相同数据重复推送", "不重复创建，返回成功并标记 duplicate: true。"),
            ("仓库尚未开始收货，预报发生变化", "更新最新预报，标记 updated: true，并保留修改记录。"),
            ("仓库已经开始收货，预报发生变化", "返回 AIR_SHIPMENT_LOCKED，不修改预报数据。"),
        ],
        [3400, 5960],
        font_size=9.0,
    )

    doc.add_page_break()
    add_heading(doc, "五、接口二：统一逐票面单推送", 1)
    add_code_block(doc, "POST /api/v1/label-pushes")
    add_body(doc, "每次请求只推送一票包裹及一份 PDF 面单。本接口统一处理首次面单、迟到补传、同转单号换 PDF 和转单号变更，不再设置独立换面单接口。CM-HUB 在单号关系和 PDF 均可靠保存后返回 HTTP 200。")

    add_heading(doc, "5.1 请求字段", 2)
    add_table(
        doc,
        ["字段", "类型", "必填", "规则"],
        [
            ("airWaybillNo", "string", "是", "已成功创建的空运提单号。"),
            ("originalTrackingNo", "string", "是", "原单号；去除首尾空格后最长 128 字符。"),
            ("transferTrackingNo", "string", "是", "当前转单号/快递单号；最长 128 字符。"),
            ("labelBase64", "string", "是", "PDF 原始字节的标准 Base64，不携带 data:application/pdf;base64, 前缀。"),
            ("replacementReason", "string", "否", "补传或更新原因，最长 200 字符；未提供时由系统记录默认原因。"),
        ],
        [2300, 1000, 700, 5360],
        font_size=8.8,
    )

    add_heading(doc, "5.2 请求示例", 2)
    add_code_block(doc, "{\n  \"airWaybillNo\": \"180-98109734\",\n  \"originalTrackingNo\": \"HHWV06218005702YQ\",\n  \"transferTrackingNo\": \"9400111899560000000000\",\n  \"labelBase64\": \"JVBERi0xLjcKJc...（完整 PDF Base64）\",\n  \"replacementReason\": \"转单号变更后重新生成面单\"\n}")

    add_heading(doc, "5.3 成功响应", 2)
    add_code_block(doc, "{\n  \"code\": \"SUCCESS\",\n  \"message\": \"接收成功\",\n  \"data\": {\n    \"airWaybillNo\": \"180-98109734\",\n    \"originalTrackingNo\": \"HHWV06218005702YQ\",\n    \"transferTrackingNo\": \"9400111899560000000000\",\n    \"operation\": \"CREATED\",\n    \"labelVersion\": 1,\n    \"duplicate\": false,\n    \"latePush\": false,\n    \"relationshipChanged\": false,\n    \"reprintRequired\": false\n  },\n  \"requestId\": \"TYG-20260831-000201\"\n}")
    add_table(
        doc,
        ["operation", "含义"],
        [
            ("CREATED", "首次创建包裹和面单。"),
            ("DUPLICATE", "单号关系和 PDF 均相同，返回原结果，不创建新版本。"),
            ("PDF_REPLACED", "转单号不变、PDF 变化，创建新版本并设为当前有效。"),
            ("TRACKING_AND_PDF_UPDATED", "原单号不变、转单号变化，原子更新关系和面单版本。"),
            ("FILE_RESTORED", "PDF 已按保存策略删除，TYG 重传后恢复文件并重新计算 7 天。"),
        ],
        [3100, 6260],
        font_size=8.8,
    )

    add_heading(doc, "5.4 PDF 与报文限制", 2)
    add_table(
        doc,
        ["项目", "限制"],
        [
            ("文件格式", "仅支持 PDF；不接受 JPG、PNG、HTML 或第三方下载链接。"),
            ("Base64", "使用标准 Base64，不含 data:application/pdf;base64, 前缀和换行。"),
            ("原始 PDF 大小", "最大 5 MiB。"),
            ("完整 JSON 请求体", "最大 7 MiB。"),
            ("服务端校验", "解码后检查 PDF 完整性，并计算 SHA-256 用于去重和审计。"),
        ],
        [2500, 6860],
        font_size=9.0,
    )

    add_heading(doc, "5.5 首次、重复、补传与覆盖规则", 2)
    add_table(
        doc,
        ["情形", "系统处理"],
        [
            ("提单不存在", "返回 AIR_SHIPMENT_NOT_FOUND，不暂存孤立面单。"),
            ("原单号不存在", "创建包裹、转单号关系和面单版本 1。"),
            ("原单号、转单号和 PDF 均相同", "按重复请求返回成功，不重复保存或计数。"),
            ("原单号和转单号相同、PDF 不同", "创建新 PDF 版本并设为当前有效面单。"),
            ("原单号相同，原 A→B 现改为 A→C", "若 C 未绑定其他原单号，原子更新为 A→C 并创建新面单版本；B 和旧 PDF 失效但保留审计。"),
            ("新转单号已绑定其他原单号", "返回 TRACKING_ALREADY_BOUND，不覆盖其他包裹。"),
            ("原单号已属于另一张空运提单", "返回 TRACKING_ALREADY_BOUND，不自动改变提单归属。"),
            ("提单已到仓或已关闭", "仍正常接收首次、补传和更新；返回 latePush: true。"),
        ],
        [3550, 5810],
        font_size=8.65,
    )

    add_heading(doc, "5.6 原子更新、打印与版本", 2)
    version_rules = [
        "转单号关系和 PDF 新版本必须作为一个原子操作完成；任一步失败时，旧转单号和旧 PDF 继续有效。",
        "更新成功后只将最新版本用于扫描和后续打印；旧转单号、旧 PDF、哈希、接收时间和操作原因保留 2 年审计。",
        "旧面单已经打印时，CM-HUB 设置 reprintRequired: true，并提醒仓库重新打印。",
        "更换 PDF 或转单号不增加有效包裹数，只增加面单版本数。",
        "每个 PDF 版本分别从成功接收时间计算 7×24 小时保存期。",
        "同一 originalTrackingNo 的业务变更必须由 TYG 按顺序推送，并等待前一版本成功后再提交下一版本。",
    ]
    for rule in version_rules:
        add_bullet(doc, rule)

    add_heading(doc, "5.7 包裹计数与到仓差异", 2)
    add_body(doc, "仓库实收数量与面单接收数量是两个独立指标。例如预报 5,000、实收到仓 5,000、已收到面单 4,500 时，系统显示实收 5,000、面单待补 500；接口不会因到仓或关闭而停止接收剩余 500 票。")
    add_table(
        doc,
        ["唯一原单号数", "相对 forecastPackages", "显示状态"],
        [
            ("小于", "尚未收齐", "显示已收面单、待补数量和完成比例。"),
            ("等于", "一致", "显示面单已收齐。"),
            ("大于", "超过预报", "继续接收并产生差异提醒，不丢失有效面单。"),
        ],
        [2500, 2500, 4360],
        font_size=9.0,
    )

    doc.add_page_break()
    add_heading(doc, "六、可靠保存、幂等、重试与并发", 1)
    add_heading(doc, "6.1 成功语义", 2)
    add_callout(doc, "可靠接收", "只有在业务数据和 PDF 面单均已可靠保存后，CM-HUB 才返回 HTTP 200 和 code: SUCCESS。成功不代表仓库已打印，但代表该票数据不会因后续异步流程而丢失。", fill=LIGHT_BLUE, title_color=NAVY)

    add_heading(doc, "6.2 初始容量约定", 2)
    add_table(
        doc,
        ["项目", "第一阶段约定"],
        [
            ("限流", "每分钟 1,200 次请求，以最终联调配置为准。"),
            ("建议并发", "TYG 客户端不超过 20 个并发请求。"),
            ("集中推送", "5 万票按初始限流约需 42 分钟；具体以真实面单压测为准。"),
            ("超限处理", "返回 HTTP 429，并通过 Retry-After 指示等待时间。"),
            ("后续目标", "持续约 100 票/秒，5 万票在 10 分钟内可靠接收；须经真实 PDF 压测通过后再作为正式承诺。"),
        ],
        [2500, 6860],
        font_size=9.0,
    )

    add_heading(doc, "6.3 幂等与自动重试", 2)
    add_body(doc, "每个真实业务版本使用新的 Idempotency-Key；网络超时后的重试必须复用原键和完全相同的请求体。同一原单号不得并发提交多个不同版本，也不得将旧请求改用新键重发，否则第一阶段无法判断业务先后。")
    add_body(doc, "网络超时、HTTP 429、500 或 503 可以使用完全相同的请求内容重试。建议最多重试 5 次，等待间隔为 1 秒、2 秒、4 秒、8 秒和 16 秒，并加入随机抖动。")
    add_body(doc, "字段错误、PDF 错误、文件过大、单号冲突和预报锁定不应原样自动重试。连续失败 5 次后应停止自动重试并告警。")

    doc.add_page_break()
    add_heading(doc, "七、数据保存与隐私", 1)
    add_table(
        doc,
        ["数据", "保存期限", "到期处理"],
        [
            ("每个版本的 PDF 原文件", "自成功接收起 7×24 小时", "自动删除，删除后不能查看或重新打印。"),
            ("完整 labelBase64", "不长期保存", "解码和校验完成后不作为业务数据保留。"),
            ("空运提单号、原单号、转单号", "2 年", "按数据保留策略到期处理。"),
            ("PDF 哈希、接收时间、版本与审计记录", "2 年", "用于对账、重复识别和操作审计。"),
        ],
        [3000, 2450, 3910],
        font_size=8.9,
    )
    add_body(doc, "PDF 保存期内可在 CM-HUB 仓库内部系统查看和重新打印；本期不向 TYG 提供状态查询或 PDF 下载接口。PDF 删除后，如 TYG 重新推送完全相同的数据，系统可重新保存 PDF 并重新计算 7 天保存期。")
    add_bullet(doc, "接口日志不得记录完整 API Key、完整 Base64 或面单中的个人信息。")
    add_bullet(doc, "PDF 应保存于私有存储，并执行最小权限访问和操作审计。")
    add_bullet(doc, "API Key 泄露后应立即停用并换发。")

    add_heading(doc, "八、错误码", 1)
    add_table(
        doc,
        ["HTTP", "错误码", "含义与处理"],
        [
            ("400", "VALIDATION_ERROR", "字段缺失、类型或长度错误；修正后再提交。"),
            ("400", "INVALID_BASE64", "labelBase64 无法解码；修正编码。"),
            ("401", "INVALID_API_KEY", "API Key 无效、已停用或环境不匹配。"),
            ("404", "AIR_SHIPMENT_NOT_FOUND", "空运提单不存在；先创建提单预报。"),
            ("409", "AIR_SHIPMENT_LOCKED", "仓库已开始收货，不能修改预报。"),
            ("400", "IDEMPOTENCY_KEY_REQUIRED", "缺少幂等键；补充后重新提交。"),
            ("409", "IDEMPOTENCY_CONFLICT", "同一幂等键对应不同请求体；不得自动重试。"),
            ("409", "TRACKING_ALREADY_BOUND", "原单号跨提单冲突，或转单号已绑定其他原单号。"),
            ("413", "PAYLOAD_TOO_LARGE", "PDF 或 JSON 请求体超过限制。"),
            ("415", "UNSUPPORTED_MEDIA_TYPE", "Content-Type 不是 application/json。"),
            ("422", "INVALID_LABEL_PDF", "Base64 可解码，但内容不是完整 PDF。"),
            ("429", "RATE_LIMITED", "降低并发并按 Retry-After 重试。"),
            ("500/503", "INTERNAL_ERROR", "我方临时错误，按重试规则提交相同内容。"),
        ],
        [700, 3100, 5560],
        font_size=8.55,
    )

    add_heading(doc, "九、首轮联调验收", 1)
    checklist = [
        "TYG 使用测试 API Key 成功创建一张含预报箱数、包裹数、重量和单位的空运提单。",
        "重复推送完全相同的提单预报，确认不产生重复记录。",
        "在仓库未收货时更新预报，并验证修改审计；模拟开始收货后验证预报锁定。",
        "为该提单推送不少于 10 票真实格式测试面单，并验证可靠保存和重复推送。",
        "验证提单不存在、非法 Base64、非 PDF、超尺寸文件和单号冲突错误。",
        "验证提单到仓或关闭后仍可首次补传面单，并显示关闭后补推标识。",
        "使用同一面单接口更换 PDF，确认生成新版本且仓库扫描和打印使用最新有效 PDF。",
        "将原 A→B 更新为 A→C，确认关系和 PDF 原子生效，旧 B 失效并保留审计；C 已占用时必须拒绝。",
        "验证同一幂等键重放返回原结果、同键不同报文返回冲突，旧版本重试不会覆盖新版本。",
        "验证 PDF 7 天删除策略及单号、哈希和审计记录 2 年保留策略。",
        "使用 TYG 真实 PDF 样本进行集中推送压测，并确认 429、Retry-After 和自动重试行为。",
        "双方确认测试通过后，另行交付生产 API Key 并安排上线窗口。",
    ]
    for item in checklist:
        add_number(doc, item)

    add_heading(doc, "十、TYG 联调前需提供", 1)
    add_bullet(doc, "几份真实 PDF 面单样本，包括常规大小和最大文件样本。")
    add_bullet(doc, "单次集中推送的最大提单数、最大包裹数和预计推送时间窗口。")
    add_bullet(doc, "空运提单号、箱数、包裹数、重量和单位在 TYG 系统中的实际字段名称。")
    add_bullet(doc, "TYG 出站服务器公网 IP（如需配置 IP 白名单）。")
    add_bullet(doc, "技术联系人、联调时间窗口和故障通知方式。")

    props = doc.core_properties
    props.title = "CM-HUB × TYG 系统对接 API 接口文档 v1.1"
    props.subject = "空运提单预报与统一逐票面单推送"
    props.author = "CM-HUB"
    props.keywords = "TYG, API, 空运提单, PDF, Base64, 面单, 版本化覆盖"
    doc.save(OUT)


if __name__ == "__main__":
    build()
